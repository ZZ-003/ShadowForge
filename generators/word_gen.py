from docx import Document
from docx.shared import Pt, Inches

class WordGenerator:
    def __init__(self, font_name="Calibri", font_size=11):
        self.font_name = font_name
        self.font_size = font_size

    def generate_docx(self, text, output_path, secret=None, secret_placeholder="SECRET_HERE"):
        """
        Generates a Word document from text.
        """
        if secret and secret_placeholder in text:
            final_text = text.replace(secret_placeholder, secret)
        else:
            final_text = text
            
        try:
            document = Document()
            
            # Simple markdown-ish parsing
            lines = final_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                if line.startswith('# '):
                    # H1
                    document.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    # H2
                    document.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # H3
                    document.add_heading(line[4:], level=3)
                else:
                    # Normal
                    p = document.add_paragraph(line)
                    p.style = document.styles['Normal']
                    
            document.save(output_path)
            return output_path
            
        except Exception as e:
            print(f"Error generating Word doc: {e}")
            return None
